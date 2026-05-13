from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ACCENT = "2F5D62"
LIGHT = "E7F0EF"
BORDER = "B8C8C6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_as_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = width_to_inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))


def width_to_inches(dxa):
    return Inches(dxa / 1440)


def apply_styles(doc, title, landscape=False):
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        content_width = 14250
    else:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        content_width = 9060

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = title
    header.style = doc.styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(95, 108, 110)

    footer = section.footer.paragraphs[0]
    footer.text = "Autoostas informācijas sistēma"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(95, 108, 110)
    return content_width


def add_metadata(doc, rows, table_width=9060):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    first_width = 2100
    set_table_width(table, [first_width, table_width - first_width])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT)
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
        cells[0].paragraphs[0].runs[0].bold = True


def parse_table(lines, start):
    table_lines = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def add_markdown_table(doc, rows, table_width=9060):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    weights = [1] * col_count
    if col_count >= 3:
        weights = [0.8] + [1.5] * (col_count - 2) + [1.2]
    total = table_width
    width_sum = sum(weights)
    widths = [int(total * weight / width_sum) for weight in weights]
    widths[-1] += total - sum(widths)

    for row_index, row in enumerate(rows):
        table_row = table.add_row()
        set_row_cant_split(table_row)
        if row_index == 0:
            set_row_as_table_header(table_row)
        cells = table_row.cells
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            cells[col_index].text = text.replace("`", "")
            if row_index == 0:
                set_cell_shading(cells[col_index], LIGHT)
            set_cell_border(cells[col_index])
            set_cell_margins(cells[col_index])
            cells[col_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[col_index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.8 if col_count > 4 else 9.4)
                    if row_index == 0:
                        run.bold = True
    set_table_width(table, widths)
    doc.add_paragraph()


def add_code_block(doc, code_lines, table_width=9060):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [table_width])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F6F8FA")
    set_cell_border(cell, "D5DDE0")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.4)
    doc.add_paragraph()


def add_markdown(doc, markdown, table_width=9060):
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            add_code_block(doc, code, table_width=table_width)
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows, table_width=table_width)
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            index += 1
            continue
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(stripped[2:], style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(numbered.group(1), style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            index += 1
            continue
        paragraph = doc.add_paragraph(stripped.replace("`", ""))
        index += 1


def build_docx(source_name, output_name, title, subtitle, landscape=False):
    doc = Document()
    table_width = apply_styles(doc, title, landscape=landscape)
    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.add_run(title)
    subtitle_paragraph = doc.add_paragraph(subtitle, style="Subtitle")
    subtitle_paragraph.runs[0].font.name = "Arial"
    subtitle_paragraph.runs[0].font.size = Pt(11)
    subtitle_paragraph.runs[0].font.color.rgb = RGBColor(95, 108, 110)
    add_metadata(
        doc,
        [
            ("Projekts", "Autoostas informācijas sistēma"),
            ("Datums", "2026-05-12"),
            ("Tehnoloģijas", "HTML, CSS, JavaScript, Python, SQLite"),
        ],
        table_width=table_width,
    )
    doc.add_paragraph()
    source = (DOCS / source_name).read_text(encoding="utf-8")
    add_markdown(doc, source, table_width=table_width)
    output = DOCS / output_name
    doc.save(output)
    return output


def main():
    outputs = [
        build_docx(
            "autoostas-sistemas-apraksts.md",
            "Autoostas_sistemas_apraksts.docx",
            "Autoostas sistēmas apraksts",
            "Īss prasību apraksts pēc dotā parauga struktūras.",
        ),
        build_docx(
            "lietotaja-interfeisa-wireframe.md",
            "Autoostas_lietotaja_interfeisa_wireframe.docx",
            "Autoostas lietotāja interfeisa wireframe",
            "Galveno ekrānu izvietojuma plāns.",
        ),
        build_docx(
            "projekta-dokumentacija.md",
            "Autoostas_projekta_dokumentacija.docx",
            "Autoostas projekta dokumentācija",
            "Prasības, modeļi, datu bāze, interfeiss, implementācija un testēšana.",
        ),
        build_docx(
            "ievades-datu-apraksts.md",
            "Autoostas_ievades_datu_apraksts.docx",
            "Autoostas ievades datu apraksts",
            "Atļautie un aizliegtie dati, pārbaudes, saglabāšanas vietas un datu izmantošanas mērķi.",
            landscape=True,
        ),
        build_docx(
            "iesniegsanas-parbaude-un-melnraksti.md",
            "Autoostas_iesniegsanas_parbaude_un_melnraksti.docx",
            "Autoostas iesniegšanas pārbaude un melnraksti",
            "Modeļi, prasības, testēšanas tabulas, ekvivalences klases, relācijas modelis un wireframe.",
            landscape=True,
        ),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
